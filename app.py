import streamlit as st
import pandas as pd
import numpy as np
import joblib
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
import google.generativeai as genai

# --------------------------
# Load model, encoders, scaler
# --------------------------
rf_model = joblib.load("random_forest_model.pkl")
label_encoders = joblib.load("label_encoders.pkl")
scaler = joblib.load("scaler.pkl")

# --------------------------
# Configure Gemini API
# --------------------------
genai.configure(api_key="AIzaSyC5XhnlGY0Ie1G045hrJo17V37Nr0chsd4")

# --------------------------
# Streamlit UI
# --------------------------
st.set_page_config(page_title="E-commerce Churn Prediction", layout="wide")
st.title("🛒 E-commerce Customer Churn Prediction")
st.write("Enter customer details to predict churn and get brief actionable tips.")

# Input form
with st.form("input_form"):
    order_count = st.number_input("Order Count", min_value=0, value=5)
    total_spent = st.number_input("Total Spent", min_value=0.0, value=10000.0)
    last_month_expenditure = st.number_input("Last Month Expenditure", min_value=0.0, value=2000.0)
    coupon_used = st.selectbox("Coupon Used", ["Yes", "No"])
    preferred_category = st.selectbox(
        "Preferred Category",
        ["Apparel", "Books", "Electronics", "Food", "Groceries"]
    )
    payment_method = st.selectbox("Payment Method", ["COD", "Credit Card", "UPI", "Wallet"])
    location_type = st.selectbox("Location Type", ["Tier 1", "Tier 2", "Tier 3"])
    rating_last_purchase = st.slider("Rating of Last Purchase", 1, 5, 3)

    submit_btn = st.form_submit_button("Predict Churn")

if submit_btn:
    # Create DataFrame from inputs
    input_df = pd.DataFrame([{
        "order_count": order_count,
        "total_spent": total_spent,
        "last_month_expenditure": last_month_expenditure,
        "coupon_used": coupon_used,
        "preferred_category": preferred_category,
        "payment_method": payment_method,
        "location_type": location_type,
        "rating_last_purchase": rating_last_purchase
    }])

    # Encode categorical features
    for col in ["coupon_used", "preferred_category", "payment_method", "location_type"]:
        le = label_encoders[col]
        input_df[col] = le.transform(input_df[col])

    # Feature engineering
    input_df['avg_order_value'] = input_df.apply(
        lambda row: row['total_spent'] / row['order_count'] if row['order_count'] > 0 else 0,
        axis=1
    )
    input_df['spend_change'] = input_df.apply(
        lambda row: row['last_month_expenditure'] / (row['total_spent'] - row['last_month_expenditure'])
        if (row['total_spent'] - row['last_month_expenditure']) > 0 else 0,
        axis=1
    )
    input_df['is_high_rating'] = (input_df['rating_last_purchase'] >= 4).astype(int)
    cat_freq = label_encoders["preferred_category"].classes_
    category_encoded_map = dict(zip(range(len(cat_freq)), np.linspace(1, 0, len(cat_freq))))
    input_df['category_encoded'] = input_df['preferred_category'].map(category_encoded_map)
    input_df['no_purchase_last_month'] = (input_df['last_month_expenditure'] == 0).astype(int)
    input_df['order_frequency_score'] = input_df.apply(
        lambda row: row['order_count'] / (row['total_spent'] / 1000) if row['total_spent'] > 0 else 0,
        axis=1
    )
    input_df['spend_trend'] = input_df.apply(
        lambda row: (row['last_month_expenditure'] / row['avg_order_value']) if row['avg_order_value'] > 0 else 0,
        axis=1
    )

    # Scale numeric columns
    cols_to_scale = [
        'order_count', 'total_spent', 'last_month_expenditure',
        'rating_last_purchase', 'avg_order_value', 'spend_change',
        'category_encoded', 'order_frequency_score', 'spend_trend'
    ]
    input_df[cols_to_scale] = scaler.transform(input_df[cols_to_scale])

    # Predict churn
    churn_class = rf_model.predict(input_df)[0]
    churn_pred = "Likely to Churn" if churn_class == 1 else "Not Likely to Churn"

    st.subheader(f"Prediction: {churn_pred}")

    # Gemini AI Explanation (brief version)
    prompt = f"""
    You are a business consultant for e-commerce.
    Customer preferred category: {preferred_category}.
    Prediction: {churn_pred}.
    Provide:
    1. 2-3 brief reasons for churn/non-churn in simple words.
    2. 3 brief tips and suggestions personalized for this preferred category.
    Keep it concise for non-technical people.
    """
    model = genai.GenerativeModel("gemini-1.5-flash")
    try:
        gemini_response = model.generate_content(prompt)
        tips_text = gemini_response.text.strip()
        st.subheader("📋 Tips & Insights")
        st.write(tips_text)
    except Exception as e:
        tips_text = "Gemini API quota exceeded. Please try again later."
        st.write(tips_text)

    # Visual Insights
    st.subheader("📊 Visual Insights")
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.bar(
        ["Total Spent", "Last Month Expenditure", "Avg Order Value"],
        [total_spent, last_month_expenditure, input_df['avg_order_value'].iloc[0]],
        color=['#1f77b4', '#ff7f0e', '#2ca02c']
    )
    ax.set_ylabel("Amount")
    ax.set_title("Customer Spending Overview")
    st.pyplot(fig)

    # Enhanced line chart for multiple features
    st.subheader("📈 Comparison Line Chart")
    metrics = ['Total Spent', 'Last Month Expense', 'Avg Order Value', 'Spend Trend', 'Order Frequency']
    values = [
        total_spent,
        last_month_expenditure,
        input_df['avg_order_value'].iloc[0],
        input_df['spend_trend'].iloc[0],
        input_df['order_frequency_score'].iloc[0]
    ]
    fig2, ax2 = plt.subplots(figsize=(8, 4))
    ax2.plot(metrics, values, marker='o', linestyle='-', color='DarkGreen', linewidth=2)
    ax2.set_ylabel("Value")
    ax2.set_title("Customer Metrics Comparison")
    st.pyplot(fig2)

    # Save report to Word including charts
    def create_word_doc():
        doc = Document()
        doc.add_heading("E-commerce Churn Prediction Report", 0)
        doc.add_paragraph(f"Prediction: {churn_pred}")

        doc.add_heading("Tips & Insights", level=1)
        doc.add_paragraph(tips_text)

        doc.add_heading("Metrics", level=2)
        for metric, val in zip(metrics, values):
            doc.add_paragraph(f"{metric}: {val}")

        # Add bar chart
        fig_buffer = BytesIO()
        fig.savefig(fig_buffer, format='png', bbox_inches='tight')
        fig_buffer.seek(0)
        doc.add_heading("Customer Spending Overview", level=2)
        doc.add_picture(fig_buffer, width=Inches(5))

        # Add line chart
        fig2_buffer = BytesIO()
        fig2.savefig(fig2_buffer, format='png', bbox_inches='tight')
        fig2_buffer.seek(0)
        doc.add_heading("Customer Metrics Comparison", level=2)
        doc.add_picture(fig2_buffer, width=Inches(5))

        # Save to BytesIO
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer


    # Create Word doc buffer
    doc_buffer = create_word_doc()

    # Directly use download_button
    st.download_button(
        label="📥 Download Report as Word",
        data=doc_buffer,
        file_name="churn_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

